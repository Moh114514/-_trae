# -*- coding: utf-8 -*-
"""
本地知识库系统
使用 ChromaDB + sentence-transformers 实现 RAG
支持格式：.md, .txt, .pdf, .docx, .doc
"""
import os
import sys
import hashlib

# 设置编码
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# 设置 Hugging Face 镜像（解决网络问题）
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

from chromadb import Client
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import chromadb

# PDF 支持
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# DOCX 支持
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# 支持的文件格式
SUPPORTED_FORMATS = ['.md', '.txt', '.pdf', '.docx', '.doc']

# 配置
KNOWLEDGE_DIR = r"E:\openclaw-project\workspace\Fuwu\knowledge_docs"  # 知识库文档目录
DB_PATH = r"E:\openclaw-project\workspace\Fuwu\chroma_db"  # 向量数据库目录

class KnowledgeBase:
    def __init__(self):
        print("初始化知识库...")
        
        # 初始化嵌入模型
        print("加载嵌入模型...")
        self.embedder = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
        
        # 初始化 ChromaDB
        print("初始化向量数据库...")
        self.client = chromadb.PersistentClient(path=DB_PATH)
        self.collection = self.client.get_or_create_collection(
            name="fuwu_knowledge",
            metadata={"description": "建筑能源管理知识库"}
        )
        
        print("知识库初始化完成！")
    
    def load_documents(self, rebuild=False):
        """
        加载所有文档（支持热更新）
        
        Args:
            rebuild: True = 清空后重建，False = 增量更新（检测文件变化）
        """
        print("\n加载文档...")
        
        # 显示支持的格式
        print(f"  支持格式: {', '.join(SUPPORTED_FORMATS)}")
        if not PDF_SUPPORT:
            print("  ⚠️ PDF支持未安装 (pip install pdfplumber)")
        if not DOCX_SUPPORT:
            print("  ⚠️ DOCX支持未安装 (pip install python-docx)")
        
        # 获取已有文档信息
        existing_files = {}  # {filename: {"hash": hash, "ids": [id1, id2, ...]}}
        if not rebuild:
            try:
                existing = self.collection.get()
                # 构建文件级别的索引
                for i, meta in enumerate(existing['metadatas']):
                    filename = meta.get('source', '')
                    file_hash = meta.get('file_hash', '')
                    doc_id = existing['ids'][i]
                    
                    if filename not in existing_files:
                        existing_files[filename] = {"hash": file_hash, "ids": []}
                    existing_files[filename]["ids"].append(doc_id)
                    # 使用第一个块的 hash 作为文件 hash
                    if file_hash and not existing_files[filename]["hash"]:
                        existing_files[filename]["hash"] = file_hash
                
                print(f"  已有 {len(existing['ids'])} 个文档块，来自 {len(existing_files)} 个文件")
            except Exception as e:
                print(f"  获取已有文档失败: {e}")
        
        documents = []
        metadatas = []
        ids = []
        deleted_count = 0
        updated_count = 0
        new_count = 0
        skipped_count = 0
        categories = {}
        
        # 记录处理过的文件
        processed_files = set()
        
        # 递归遍历知识库目录
        for root, dirs, files in os.walk(KNOWLEDGE_DIR):
            # 计算相对路径作为分类
            rel_path = os.path.relpath(root, KNOWLEDGE_DIR)
            category = rel_path if rel_path != "." else "根目录"
            
            for filename in files:
                ext = os.path.splitext(filename)[1].lower()
                if ext not in SUPPORTED_FORMATS:
                    continue
                
                filepath = os.path.join(root, filename)
                processed_files.add(filename)
                
                try:
                    # 计算文件哈希
                    file_hash = self._get_file_hash(filepath)
                    
                    # 检查文件是否有变化
                    if not rebuild and filename in existing_files:
                        old_hash = existing_files[filename].get("hash", "")
                        
                        if old_hash == file_hash:
                            # 文件未变化，跳过
                            skipped_count += len(existing_files[filename]["ids"])
                            continue
                        else:
                            # 文件已变化，删除旧块
                            old_ids = existing_files[filename]["ids"]
                            self.collection.delete(ids=old_ids)
                            deleted_count += len(old_ids)
                            print(f"  [{category}] 更新 {filename} (删除 {len(old_ids)} 个旧块)")
                            updated_count += 1
                    
                    # 读取文件内容
                    content = self._read_file(filepath, ext)
                    if not content:
                        print(f"  跳过 {filename}: 内容为空")
                        continue
                    
                    # 分割文档为多个块
                    chunks = self._split_text(content, chunk_size=500)
                    
                    file_new_chunks = 0
                    for i, chunk in enumerate(chunks):
                        # 使用相对路径作为文档ID，确保唯一性
                        doc_id = f"{rel_path}_{filename}_{i}".replace("\\", "/")
                        
                        documents.append(chunk)
                        # 元数据包含分类信息和文件哈希
                        metadatas.append({
                            "source": filename,
                            "category": category,
                            "path": filepath,
                            "chunk": i,
                            "file_hash": file_hash  # 存储文件哈希用于检测变化
                        })
                        ids.append(doc_id)
                        file_new_chunks += 1
                    
                    # 统计分类
                    if category not in categories:
                        categories[category] = {"files": 0, "chunks": 0}
                    categories[category]["files"] += 1
                    categories[category]["chunks"] += file_new_chunks
                    
                    if filename not in existing_files:
                        print(f"  [{category}] 新增 {filename}: {file_new_chunks} 个块")
                        new_count += 1
                        
                except Exception as e:
                    print(f"  加载 {filename} 失败: {e}")
        
        # 检测已删除的文件（知识库中存在但目录中不存在）
        if not rebuild:
            deleted_files = set(existing_files.keys()) - processed_files
            for filename in deleted_files:
                old_ids = existing_files[filename]["ids"]
                self.collection.delete(ids=old_ids)
                deleted_count += len(old_ids)
                print(f"  删除不存在的文件: {filename} ({len(old_ids)} 个块)")
        
        # 添加到向量数据库
        if documents:
            print(f"\n添加 {len(documents)} 个文档块到向量数据库...")
            self.collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            print("文档添加完成！")
            
            # 显示分类统计
            print("\n分类统计：")
            for cat, stats in sorted(categories.items()):
                if stats["chunks"] > 0:
                    print(f"  {cat}: {stats['files']} 个文件, {stats['chunks']} 个块")
        
        # 显示更新摘要
        print(f"\n更新摘要：")
        print(f"  新增文件: {new_count}")
        print(f"  更新文件: {updated_count}")
        print(f"  删除块数: {deleted_count}")
        print(f"  跳过块数: {skipped_count}")
        
        return len(documents)
    
    def _get_file_hash(self, filepath):
        """计算文件的 MD5 哈希"""
        hasher = hashlib.md5()
        try:
            with open(filepath, 'rb') as f:
                # 分块读取大文件
                for chunk in iter(lambda: f.read(8192), b''):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            # 如果无法读取，使用修改时间作为后备
            mtime = os.path.getmtime(filepath)
            return f"mtime_{int(mtime)}"
    
    def _read_file(self, filepath, ext):
        """读取不同格式的文件"""
        if ext in ['.md', '.txt']:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif ext == '.pdf':
            if not PDF_SUPPORT:
                raise ImportError("需要安装 pdfplumber: pip install pdfplumber")
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        
        elif ext in ['.docx', '.doc']:
            if not DOCX_SUPPORT:
                raise ImportError("需要安装 python-docx: pip install python-docx")
            doc = Document(filepath)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            # 读取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        
        return ""
    
    def _split_text(self, text, chunk_size=500):
        """分割文本为多个块"""
        chunks = []
        lines = text.split('\n')
        current_chunk = ""
        
        for line in lines:
            if len(current_chunk) + len(line) > chunk_size:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = line + "\n"
            else:
                current_chunk += line + "\n"
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def search(self, query, n_results=3, category=None):
        """
        搜索相关文档
        
        Args:
            query: 查询文本
            n_results: 返回结果数量
            category: 可选，按分类筛选（如 "设备手册"、"运维指南"）
        """
        if category:
            # 按分类筛选搜索
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results,
                where={"category": category}
            )
        else:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results
            )
        
        return results
    
    def get_context(self, query, n_results=3, category=None):
        """
        获取查询上下文
        
        Args:
            query: 查询文本
            n_results: 返回结果数量
            category: 可选，按分类筛选
        """
        results = self.search(query, n_results, category)
        
        if results['documents']:
            context = "\n\n---\n\n".join(results['documents'][0])
            sources = [(m.get('source', '未知'), m.get('category', '未分类')) for m in results['metadatas'][0]]
            return context, sources
        
        return "", []
    
    def list_categories(self):
        """列出所有分类"""
        try:
            results = self.collection.get()
            categories = set()
            for m in results['metadatas']:
                categories.add(m.get('category', '未分类'))
            return sorted(list(categories))
        except:
            return []
    
    def get_stats(self):
        """获取知识库统计"""
        count = self.collection.count()
        categories = self.list_categories()
        return {
            "total_chunks": count,
            "categories": categories,
            "category_count": len(categories)
        }


def main():
    print("=" * 50)
    print("建筑能源知识库系统")
    print("=" * 50)
    
    # 初始化知识库
    kb = KnowledgeBase()
    
    # 加载文档
    doc_count = kb.load_documents()
    
    # 显示统计
    stats = kb.get_stats()
    print(f"\n知识库统计：")
    print(f"  总文档块数：{stats['total_chunks']}")
    
    # 测试搜索
    print("\n" + "=" * 50)
    print("测试搜索功能")
    print("=" * 50)
    
    test_queries = [
        "COP是什么意思",
        "如何检测能耗异常",
        "冷冻水供水温度的正常范围是多少"
    ]
    
    for query in test_queries:
        print(f"\n查询：{query}")
        context, sources = kb.get_context(query)
        print(f"来源：{sources}")
        print(f"上下文预览：{context[:200]}...")


if __name__ == "__main__":
    main()